r"""
apply_docx_format.py — DOCX 受控文档格式应用脚本
按照 references/docx-format.md 规范，对 pandoc 生成的 DOCX 进行精确后处理。

用法:
    python scripts/apply_docx_format.py <输入.docx> [输出.docx]
    若不指定输出路径，则覆盖输入文件。
"""

import sys
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


# ─────────────────────────────────────────────
# 工具函数
# ─────────────────────────────────────────────

def set_font(run, cn_font: str, en_font: str, size_pt: float,
             bold: bool = False, color: RGBColor = None):
    """设置 run 的中西文字体、字号、加粗、颜色"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    rPr = run._r.get_or_add_rPr()

    # 西文字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    # 中文字体
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), cn_font)


def set_para_spacing(para, before_twips: int = 0, after_twips: int = 0,
                     line_rule=WD_LINE_SPACING.SINGLE, line_val=None):
    """设置段落间距和行距"""
    pf = para.paragraph_format
    pf.space_before = Twips(before_twips)
    pf.space_after = Twips(after_twips)
    pf.line_spacing_rule = line_rule
    if line_val is not None:
        pf.line_spacing = line_val


def remove_underline(para):
    """移除段落中所有 run 的下划线"""
    for run in para.runs:
        run.font.underline = False
    # 同时清除 pPr 级别的边框（pandoc 有时会加底部边框模拟下划线）
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)


def set_cell_shading(cell, fill_hex: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)


def set_cell_borders(cell, color_hex: str = 'D9D9D9'):
    """设置单元格边框为淡灰色细线"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        existing = tcBorders.find(qn(f'w:{side}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(border)


def set_cell_vertical_align(cell, align=WD_ALIGN_VERTICAL.CENTER):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), 'center' if align == WD_ALIGN_VERTICAL.CENTER else 'top')


def guess_heading_level(para) -> int:
    """返回段落的标题级别 (1/2/3/4)，0 表示非标题"""
    style_name = para.style.name.lower()
    for lvl in [1, 2, 3, 4]:
        if f'heading {lvl}' in style_name or style_name == f'h{lvl}':
            return lvl
    # pandoc 有时用 style_name like "Heading1" 或 "heading1"
    for lvl in [1, 2, 3, 4]:
        if style_name.endswith(str(lvl)) and 'head' in style_name:
            return lvl
    return 0


# ─────────────────────────────────────────────
# 核心格式应用
# ─────────────────────────────────────────────

def apply_heading1(para):
    """大标题 H1：宋体/TNR 小二18pt 加粗 居中 段后1行"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, before_twips=0, after_twips=240,
                     line_rule=WD_LINE_SPACING.SINGLE)
    remove_underline(para)
    for run in para.runs:
        set_font(run, '宋体', 'Times New Roman', 18, bold=True,
                 color=RGBColor(0, 0, 0))


def apply_heading2(para):
    """一级标题 H2：黑体/TNR 四号14pt 加粗 左对齐 段前1行段后0.5行 无下划线"""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    set_para_spacing(para, before_twips=240, after_twips=120,
                     line_rule=WD_LINE_SPACING.SINGLE)
    remove_underline(para)
    for run in para.runs:
        set_font(run, '黑体', 'Times New Roman', 14, bold=True,
                 color=RGBColor(0, 0, 0))


def apply_heading3(para):
    """二级标题 H3：黑体/TNR 小四12pt 加粗 左对齐 段前后各0.5行"""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    set_para_spacing(para, before_twips=120, after_twips=120,
                     line_rule=WD_LINE_SPACING.SINGLE)
    remove_underline(para)
    for run in para.runs:
        set_font(run, '黑体', 'Times New Roman', 12, bold=True,
                 color=RGBColor(0, 0, 0))


def apply_heading4(para):
    """三级标题 H4：宋体/TNR 五号10.5pt 加粗 左对齐 段前0.5行"""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    set_para_spacing(para, before_twips=120, after_twips=0,
                     line_rule=WD_LINE_SPACING.SINGLE)
    remove_underline(para)
    for run in para.runs:
        set_font(run, '宋体', 'Times New Roman', 10.5, bold=True,
                 color=RGBColor(0, 0, 0))


def apply_body(para):
    """正文：宋体/TNR 五号10.5pt 两端对齐 1.5倍行距 首行缩进2字符 段后0.5行"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    # 首行缩进 2字符 ≈ 2 * 10.5pt = 21pt
    pf.first_line_indent = Pt(21)
    pf.left_indent = Pt(0)
    set_para_spacing(para, before_twips=0, after_twips=120,
                     line_rule=WD_LINE_SPACING.MULTIPLE,
                     line_val=Pt(10.5 * 1.5))  # 1.5倍
    for run in para.runs:
        # 仅处理未设置特殊颜色（红色强调）的 run
        if run.font.color.rgb not in (RGBColor(0xFF, 0, 0),):
            set_font(run, '宋体', 'Times New Roman', 10.5,
                     color=RGBColor(0, 0, 0))
        else:
            # 保留红色，只更新字体和字号
            set_font(run, '宋体', 'Times New Roman', 10.5,
                     color=run.font.color.rgb)


def apply_caption(para):
    """图注：宋体/TNR 10pt 居中 单倍行距 段后0.5行"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    set_para_spacing(para, before_twips=0, after_twips=120,
                     line_rule=WD_LINE_SPACING.SINGLE)
    for run in para.runs:
        set_font(run, '宋体', 'Times New Roman', 10,
                 color=RGBColor(0, 0, 0))


def is_caption_para(para) -> bool:
    """判断是否为图注段落"""
    style_name = para.style.name.lower()
    if 'caption' in style_name or 'figure' in style_name:
        return True
    text = para.text.strip()
    # 匹配 "图 X：" 或 "图X：" 或 "Fig." 开头
    import re
    if re.match(r'^(图\s*\d|Fig\.?\s*\d)', text):
        return True
    return False


def apply_table_format(table):
    """表格：三线表斑马纹，绿色表头，淡灰边框"""
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        is_even = (row_idx % 2 == 0)  # 0-indexed: 偶数索引=奇数行（白色）

        for cell in row.cells:
            # 背景色
            if is_header:
                set_cell_shading(cell, '00B050')
            elif is_even:
                set_cell_shading(cell, 'FFFFFF')
            else:
                set_cell_shading(cell, 'F2F2F2')

            # 边框
            set_cell_borders(cell, 'D9D9D9')
            # 垂直居中
            set_cell_vertical_align(cell)

            # 单元格内段落格式
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = para.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)
                set_para_spacing(para, before_twips=60, after_twips=60,
                                 line_rule=WD_LINE_SPACING.SINGLE)
                for run in para.runs:
                    if is_header:
                        set_font(run, '微软雅黑', '微软雅黑', 10.5, bold=True,
                                 color=RGBColor(0xFF, 0xFF, 0xFF))
                    else:
                        set_font(run, '微软雅黑', '微软雅黑', 10.5, bold=False,
                                 color=RGBColor(0, 0, 0))


def is_image_para(para) -> bool:
    """判断是否为纯图片段落"""
    p_xml = para._p.xml
    return 'w:drawing' in p_xml or 'w:pict' in p_xml


def apply_image_para(para):
    """图片段落：居中 段前后各0.5行"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, before_twips=120, after_twips=120,
                     line_rule=WD_LINE_SPACING.SINGLE)


# ─────────────────────────────────────────────
# 主处理流程
# ─────────────────────────────────────────────

def apply_format(doc: Document):
    """遍历文档所有段落和表格，应用格式规范"""
    stats = {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0,
             'body': 0, 'caption': 0, 'image': 0, 'table': 0}

    for para in doc.paragraphs:
        lvl = guess_heading_level(para)

        if lvl == 1:
            apply_heading1(para)
            stats['h1'] += 1
        elif lvl == 2:
            apply_heading2(para)
            stats['h2'] += 1
        elif lvl == 3:
            apply_heading3(para)
            stats['h3'] += 1
        elif lvl == 4:
            apply_heading4(para)
            stats['h4'] += 1
        elif is_image_para(para):
            apply_image_para(para)
            stats['image'] += 1
        elif is_caption_para(para):
            apply_caption(para)
            stats['caption'] += 1
        else:
            # 跳过空段落和只含空白的段落
            if para.text.strip():
                apply_body(para)
                stats['body'] += 1

    for table in doc.tables:
        apply_table_format(table)
        stats['table'] += 1

    return stats


def main():
    if len(sys.argv) < 2:
        print("用法: python apply_docx_format.py <输入.docx> [输出.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path

    if not os.path.exists(input_path):
        print(f"[ERROR] 文件不存在: {input_path}")
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"  DOCX 格式应用脚本 (受控文档排版规范)")
    print(f"{'='*60}")
    print(f"  输入: {input_path}")
    print(f"  输出: {output_path}")
    print(f"{'='*60}\n")

    doc = Document(input_path)
    stats = apply_format(doc)
    doc.save(output_path)

    print(f"[完成] 格式应用结果：")
    print(f"  大标题(H1):   {stats['h1']} 处")
    print(f"  一级标题(H2): {stats['h2']} 处")
    print(f"  二级标题(H3): {stats['h3']} 处")
    print(f"  三级标题(H4): {stats['h4']} 处")
    print(f"  正文段落:     {stats['body']} 处")
    print(f"  图注:         {stats['caption']} 处")
    print(f"  图片段落:     {stats['image']} 处")
    print(f"  表格:         {stats['table']} 个")
    print(f"\n  输出文件: {output_path}\n")


if __name__ == "__main__":
    main()
