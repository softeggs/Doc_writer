import sys
import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def strip_title_prefix(para):
    """剔除标题开头的数字序号，例如 '3. ', '3.1. ', '第1章 ' 等"""
    # 匹配模式：数字开头后接点或空格，或者“第X章/节”开头
    prefix_pattern = re.compile(r'^(\d+(\.\d+)*[\s\.\、]*|第[一二三四五六七八九十\d]+[章节]\s*)')
    
    full_text = para.text
    match = prefix_pattern.match(full_text)
    if match:
        prefix = match.group(0)
        remaining_len = len(prefix)
        # 逐个 run 清理，直到删除了前缀对应的长度
        for run in para.runs:
            if remaining_len <= 0:
                break
            run_text_len = len(run.text)
            if run_text_len <= remaining_len:
                remaining_len -= run_text_len
                run.text = ""
            else:
                run.text = run.text[remaining_len:]
                remaining_len = 0

def apply_edan_styles(doc: Document):
    """根据 EDAN 模板规范应用样式，并自动清理冗余编号"""
    
    h1_count = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name.lower()
        text = para.text.strip()
        
        # 识别标题级别
        lvl = 0
        if 'heading 1' in style_name or style_name == 'h1': lvl = 1
        elif 'heading 2' in style_name or style_name == 'h2': lvl = 2
        elif 'heading 3' in style_name or style_name == 'h3': lvl = 3
        elif 'heading 4' in style_name or style_name == 'h4': lvl = 4
        
        if lvl > 0:
            # 对于所有标题，应用样式前先剥离 HTML 自带的编号
            strip_title_prefix(para)
            
            if lvl == 1:
                h1_count += 1
                if h1_count == 1:
                    try: para.style = doc.styles['文件名称']
                    except: para.style = doc.styles['Heading 1']
                else:
                    try: para.style = doc.styles['Heading 1']
                    except: pass
            elif lvl == 2:
                try: para.style = doc.styles['文章一级标题']
                except: pass
            elif lvl == 3:
                try: para.style = doc.styles['文章二级标题']
                except: pass
            elif lvl == 4:
                try: para.style = doc.styles['文章三级标题']
                except: pass
        
        elif 'w:drawing' in para._p.xml or 'w:pict' in para._p.xml:
            try: para.style = doc.styles['图片样式']
            except: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'caption' in style_name or re.match(r'^(图|表)\s*\d', text):
            try: para.style = doc.styles['edan题注']
            except: pass
        elif text:
            try: para.style = doc.styles['文章正文段落']
            except: pass

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    try: para.style = doc.styles['表格正文']
                    except: pass
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    if len(sys.argv) < 2:
        print("Usage: python apply_edan_format.py <input.docx> [output.docx]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path
    
    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found")
        sys.exit(1)
        
    doc = Document(input_path)
    apply_edan_styles(doc)
    doc.save(output_path)
    print(f"Successfully applied EDAN template styles and stripped prefixes in {output_path}")

if __name__ == "__main__":
    main()
